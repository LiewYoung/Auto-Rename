"""
This module provides the ReadFile class for extracting text content from various file types.

Supported file types include plain text, PDF, and DOCX. It aims to extract
a limited amount of text from the beginning of these files.
"""
from pathlib import Path
from docx import Document
import PyPDF2
from typing import Union # Union was deprecated in 3.10, use | instead

class ReadFile:
    """
    Reads and extracts text content from specified files.

    The class supports reading plain text files, PDF files (first page),
    and DOCX files (first paragraph). It limits the amount of text read
    to `max_read_size`.
    """
    def __init__(self, file_path: Union[Path, str], max_read_size: int = 1024):
        """
        Initializes the ReadFile object.

        :param file_path: Path to the file to be read (can be a Path object or string).
        :param max_read_size: Maximum number of characters to read from the file.
        """
        self.path: Union[Path, str] = file_path
        self.max_read_size: int = max_read_size
        self.extension: str = Path(self.path).suffix.lower()
        
        
    def get_file(self) -> str:
        """
        Extracts text content from the file.

        For plain text files, it reads up to `max_read_size` characters.
        For PDF files, it extracts text from the first page, up to `max_read_size`.
        For DOCX files, it extracts text from the first paragraph, up to `max_read_size`.
        If the file content is empty, or not extractable (e.g., empty PDF/DOCX),
        it returns the filename.
        If the file type is unsupported or a UnicodeDecodeError occurs,
        it returns the filename or an error message.

        :return: The extracted text content, filename, or an error message.
        """
        try:
            if self.extension not in ('.pdf','.docx'):
                # Ensure self.path is used consistently, it could be str or Path
                with open(self.path,'r', encoding='utf-8') as file: # Added encoding
                    file_include=file.read(self.max_read_size)
                    if not file_include.strip():
                        return Path(self.path).name
                    return file_include
        
            elif self.extension=='.pdf':
                with open(self.path,'rb') as file: # self.path could be str
                    reader=PyPDF2.PdfReader(file)
                    if not reader.pages:
                        return Path(self.path).name
                    page=reader.pages[0]
                    file_include=page.extract_text()
                    if not file_include or not file_include.strip():
                        return Path(self.path).name
                    return file_include[:self.max_read_size]
        
            elif self.extension=='.docx':
                doc=Document(self.path) # self.path could be str
                if not doc.paragraphs:
                    return Path(self.path).name
                text=doc.paragraphs[0].text
                if not text.strip():
                    return Path(self.path).name
                return text[:self.max_read_size]
            else:
                return f"Unsupported file type: {self.extension}"
        except UnicodeDecodeError:
            # Consider logging this error instead of printing directly if this is a library
            print(f'UnicodeDecodeError for file: {Path(self.path).name}, extension: {self.extension}')
            return Path(self.path).name
        except Exception as e: # Catch other potential errors during file processing
            print(f"Error processing file {Path(self.path).name}: {e}")
            return Path(self.path).name
        

if __name__ == '__main__':
    # Example Usage
    # Get path from command line argument or keep input for simple testing
    # For robust CLI, argparse would be better here as in chat.py
    input_path_str = input('Please input the path: ')
    # Ensure input_path is a Path object for consistency if ReadFile expects Path
    # However, __init__ now accepts Union[Path, str]
    read_file_instance = ReadFile(input_path_str)
    print(read_file_instance.get_file())
            
